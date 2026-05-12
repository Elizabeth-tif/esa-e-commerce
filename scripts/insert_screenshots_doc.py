from docx import Document
from docx.shared import Inches, Pt
import os

DOCX_IN = 'Dokumentasi_Praktikum_12_Microservices.docx'
DOCX_OUT = 'Dokumentasi_Praktikum_12_Microservices_updated.docx'
IMAGES = [
    ('postman_checkout.png', 'Postman: POST /api/v1/orders/checkout (response: started) - tangkapan layar menampilkan saga_id yang diterima saat checkout') ,
    ('postman_saga_status.png', 'Postman: GET /api/v1/orders/saga/{saga_id} (response: final state) - tangkapan layar menampilkan riwayat langkah saga dan status akhir')
]

def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style = 'Normal'

def insert_images_and_context():
    if not os.path.exists(DOCX_IN):
        print(f"Source docx not found: {DOCX_IN}")
        return 1

    doc = Document(DOCX_IN)

    add_heading(doc, 'Postman Test Screenshots & Context')
    add_paragraph(doc, 'Di bagian ini disertakan dua tangkapan layar Postman yang menunjukkan: 1) respons awal dari endpoint checkout yang mengembalikan `saga_id`, dan 2) status akhir saga (COMPLETED / FAILED / dengan fallback).')

    # Detailed context for each image
    for filename, caption in IMAGES:
        add_paragraph(doc, '')
        if os.path.exists(filename):
            try:
                doc.add_picture(filename, width=Inches(6))
                last_par = doc.paragraphs[-1]
                last_par.alignment = 0
                add_paragraph(doc, f'Keterangan: {caption}')
                print(f'Inserted image: {filename}')
            except Exception as e:
                add_paragraph(doc, f'Gagal memasukkan gambar {filename}: {e}')
                print(f'Failed to insert {filename}: {e}')
        else:
            add_paragraph(doc, f'[PLACEHOLDER] Tidak menemukan file gambar {filename} di folder kerja. Silakan unggah {filename} ke folder project dan jalankan ulang skrip ini untuk menyisipkannya.')
            add_paragraph(doc, f'Keterangan yang disarankan: {caption}')
            print(f'Placeholder added for missing image: {filename}')

    # Add more context paragraphs explaining test steps
    add_heading(doc, 'Instruksi Pengujian (Postman)')
    add_paragraph(doc, '1) POST http://localhost:3000/api/v1/orders/checkout - Body: contoh payload order (user_id, items, total_amount). Respons diharapkan 201 Created dengan JSON {"saga_id":"...","status":"started"}. Tangkap layar respons ini sebagai bukti bahwa saga telah dimulai.')
    add_paragraph(doc, '2) GET http://localhost:3000/api/v1/orders/saga/{saga_id} - Ganti {saga_id} dengan nilai dari respons checkout. Respons diharapkan menampilkan field `current_state`, `step_history`, dan `order_data`. Tangkap layar ini untuk menunjukkan apakah saga selesai normal atau menyelesaikan dengan fallback/kompensasi.')
    add_paragraph(doc, 'Catatan: Untuk mensimulasikan skenario fallback/circuit-breaker, hentikan salah satu service downstream (mis. shipping-service) lalu jalankan ulang langkah 1 untuk memicu fallback. Periksa log orchestrator untuk pesan fallback dan circuit breaker state.')

    doc.save(DOCX_OUT)
    print(f'Saved updated document as: {DOCX_OUT}')
    return 0

if __name__ == '__main__':
    exit(insert_images_and_context())
